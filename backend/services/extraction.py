import io
import pdfplumber
from docx import Document

from backend.models import ExtractedDocument, PageSpan

# ponytail: paragraph layout reconstruction skipped; .chars join is sufficient
# for offset round-trip — revisit only if a fixture fails


def extract_pdf_document(content: bytes) -> ExtractedDocument:
    """Extract PDF text plus page spans using the canonical sorted .chars join.

    Uses pdfplumber .chars sorted by (doctop, x0) — document-level Y coordinate
    keeps offsets stable across page boundaries. Do NOT call page.extract_text();
    it normalizes whitespace and corrupts offsets (Pitfall 1).
    """
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        chars = []
        for page_number, page in enumerate(pdf.pages, start=1):
            for char in page.chars:
                tagged = dict(char)
                tagged["_page_number"] = page_number
                chars.append(tagged)
    chars.sort(key=lambda c: (c["doctop"], c["x0"]))

    text_parts: list[str] = []
    page_start_offsets: dict[int, int] = {}
    current_offset = 0

    for char in chars:
        page_number = char["_page_number"]
        page_start_offsets.setdefault(page_number, current_offset)
        text = char["text"]
        text_parts.append(text)
        current_offset += len(text)

    page_spans: list[PageSpan] = []
    ordered_pages = sorted(page_start_offsets)
    for index, page_number in enumerate(ordered_pages):
        start_offset = page_start_offsets[page_number]
        end_offset = (
            page_start_offsets[ordered_pages[index + 1]]
            if index + 1 < len(ordered_pages)
            else current_offset
        )
        page_spans.append(
            PageSpan(
                page_number=page_number,
                start_offset=start_offset,
                end_offset=end_offset,
            )
        )

    return ExtractedDocument(full_text="".join(text_parts), page_spans=page_spans)


def extract_pdf(content: bytes) -> str:
    """Compatibility wrapper returning only the canonical PDF full_text."""
    return extract_pdf_document(content).full_text


def extract_docx_document(content: bytes) -> ExtractedDocument:
    """Extract DOCX text with deterministic MVP page metadata.

    ponytail: tables/text-boxes skipped (doc.paragraphs only); iterate
    doc.element.body when a table-containing fixture fails round-trip (Pitfall 5).
    """
    doc = Document(io.BytesIO(content))
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return ExtractedDocument(
        full_text=full_text,
        page_spans=[PageSpan(page_number=1, start_offset=0, end_offset=len(full_text))],
    )


def extract_docx(content: bytes) -> str:
    """Compatibility wrapper returning only the canonical DOCX full_text."""
    return extract_docx_document(content).full_text
